#!/usr/bin/env python3
"""
Create Manuscript v3 - Checklist-based template with all gaps, questions, and tasks
Designed for systematic completion of all missing elements
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# COVER PAGE
title = doc.add_paragraph()
title_run = title.add_run('MANUSCRIPT v3 - PUBLICATION CHECKLIST\n')
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(192, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Senescence Signatures and CAR-T Targets in Systemic Lupus Erythematosus\n')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta_run = meta.add_run(f'Working Document | Generated {datetime.now().strftime("%Y-%m-%d")}\n')
meta_run.font.size = Pt(10)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

status = doc.add_paragraph()
status_run = status.add_run('Status: GAPS & RESEARCH QUESTIONS IDENTIFIED | READY FOR SYSTEMATIC COMPLETION')
status_run.font.bold = True
status_run.font.color.rgb = RGBColor(192, 0, 0)
status.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# TABLE OF CONTENTS
doc.add_heading('Contents', level=1)
toc_items = [
    '1. Overview & Completion Roadmap',
    '2. Critical Research Questions (Must Answer Before Writing)',
    '3. SECTION 1: INTRODUCTION (MISSING - ~3-4 pages)',
    '4. SECTION 2: METHODS (PARTIAL - EXPAND)',
    '5. SECTION 3: RESULTS (GOOD - NEEDS REORGANIZATION)',
    '6. SECTION 4: DISCUSSION (MISSING - ~4-5 pages)',
    '7. SECTION 5: FUTURE DIRECTIONS (MISSING - ~1-2 pages)',
    '8. Literature Research Checklist',
    '9. Analysis Checklist',
    '10. Cross-Validation Checklist',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# SECTION 0: OVERVIEW
doc.add_heading('1. Overview & Completion Roadmap', level=1)

doc.add_heading('Current State', level=2)
doc.add_paragraph('Blueprint v2 = Executive Summary (good for presentations, weak for peer review)')
doc.add_paragraph('Estimated completion: 4-5 weeks of systematic work')
doc.add_paragraph()

doc.add_heading('Work Phases', level=2)
phases_table = doc.add_table(rows=1, cols=4)
phases_table.style = 'Light Grid Accent 1'
hdr = phases_table.rows[0].cells
hdr[0].text = 'Phase'
hdr[1].text = 'Duration'
hdr[2].text = 'Output'
hdr[3].text = 'Gating Criteria'

phases = [
    ('Research Sprint', '1 week', 'Answered Q1-Q10, literature review', '15+ papers read, critical gaps identified'),
    ('Methods Writing', '3 days', 'METHODS.md expanded', 'All formulas present, choices justified'),
    ('Results Reorganization', '2 days', 'RESULTS.md separated', 'Pure findings; no interpretation'),
    ('Discussion Writing', '1 week', 'DISCUSSION.md created', 'Literature integrated, mechanisms explained'),
    ('Integration', '3 days', 'Full manuscript', 'Sections flow; citations complete'),
]
for phase in phases:
    row = phases_table.add_row()
    row.cells[0].text = phase[0]
    row.cells[1].text = phase[1]
    row.cells[2].text = phase[2]
    row.cells[3].text = phase[3]

doc.add_paragraph()

# SECTION 1: CRITICAL RESEARCH QUESTIONS
doc.add_heading('2. Critical Research Questions (Must Answer Before Writing)', level=1)
doc.add_paragraph('These questions block manuscript writing. Answer them via literature search FIRST.')
doc.add_paragraph()

questions = [
    {
        'num': 'Q1',
        'question': 'SLE EPIDEMIOLOGY: Current prevalence, mortality, standard of care gaps?',
        'needed_for': 'Introduction - Problem statement',
        'search_queries': [
            'PubMed: "systemic lupus erythematosus" epidemiology 2023-2024',
            'PubMed: "lupus nephritis" incidence mortality outcomes 2022-2024',
            'Target: Find 3-5 recent epidemiology/outcomes papers',
        ],
        'acceptance_criteria': [
            'Global prevalence (millions) with recent citation',
            'Mortality rate and organ involvement statistics',
            'Current standard of care (belimumab, steroids, biologics)',
            'Why current therapies are insufficient',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q2',
        'question': 'CAR-T LANDSCAPE: CD19+ CAR-T successes in SLE? Single-target escape mechanisms?',
        'needed_for': 'Introduction - Rationale for multi-target approach',
        'search_queries': [
            'Clinical Trials.gov: CAR-T CD19 SLE lupus (active/closed trials)',
            'PubMed: "CAR-T" "CD19" SLE 2020-2024',
            'PubMed: "CAR-T escape" OR "antigen loss" mechanisms',
            'Target: Find ≥2 SLE CAR-T trials + ≥3 mechanism papers',
        ],
        'acceptance_criteria': [
            'Response rates of CD19 CAR-T in SLE trials (if exist)',
            'Alternative: Response rates in B-cell malignancies (proof of concept)',
            'Documented escape mechanisms (CD19-negative relapse)',
            'Rationale for extending beyond CD19',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q3',
        'question': 'SENESCENCE IN AUTOIMMUNITY: How does SLE senescence differ from aging senescence?',
        'needed_for': 'Introduction - Biological rationale',
        'search_queries': [
            'PubMed: "senescence" "autoimmune disease" OR "rheumatoid arthritis" 2020-2024',
            'PubMed: "p16" OR "p21" OR "SASP" lupus OR SLE 2019-2024',
            'PubMed: "senescent T cells" OR "senescent B cells" autoimmunity',
            'Target: Find 5+ papers on senescence in autoimmunity',
        ],
        'acceptance_criteria': [
            'Evidence that senescence accumulates faster in SLE vs. healthy aging',
            'SASP cytokines (IL-6, TNF, MMP) elevated in SLE',
            'Senescence-driven inflammation in lupus tissues',
            'Prior studies on p16/p21 expression in SLE patients',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q4',
        'question': 'CD38 on senescent cells: Is it a known marker or novel?',
        'needed_for': 'Discussion - Novelty claim',
        'search_queries': [
            'PubMed: "CD38" senescent cells OR senescence aging',
            'PubMed: "CD38" "T cells" OR "immune cells" expression',
            'PubMed: "daratumumab" CD38 CAR-T autoimmunity (clinical use)',
            'Target: Determine if CD38 on senescence is established or new',
        ],
        'acceptance_criteria': [
            'Prior papers describing CD38 on senescent cells (if exist)',
            'CD38 validation in CAR-T trials (solid tumors, hematologic)',
            'CD38 expression patterns in SLE vs. health',
            'Rationale: Why CD38 is relevant to senescence in SLE',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q5',
        'question': 'CD44 on senescent cells: Known marker or novel?',
        'needed_for': 'Discussion - Novelty claim',
        'search_queries': [
            'PubMed: "CD44" senescent cells aging',
            'PubMed: "CD44" "stem cell" OR "memory cells" senescence',
            'PubMed: "CD44" CAR-T trials outcomes toxicity',
            'Target: Determine if CD44 on senescence is established',
        ],
        'acceptance_criteria': [
            'Established role of CD44 as memory/stem cell marker',
            'Prior studies: CD44 on senescent cells (if any)',
            'CD44 CAR-T toxicity in preclinical/clinical (endothelium, bone marrow)',
            'Off-target risks explanation',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q6',
        'question': 'CSPG4 on immune cells: Is CSPG4 known on lymphocytes/myeloid cells?',
        'needed_for': 'Discussion - Critical novelty claim',
        'search_queries': [
            'PubMed: "CSPG4" immune cells T cells B cells macrophages',
            'PubMed: "CSPG4" senescence',
            'PubMed: "CSPG4" CAR-T trials (not just melanoma)',
            'Target: Determine if CSPG4 on SLE immune cells is novel',
        ],
        'acceptance_criteria': [
            'Literature on CSPG4 normal expression (melanoma, stromal cells)',
            'Whether CSPG4 on immune cells is established or NEW',
            'If novel: Strengthen novelty claim; if known: Cite prior work',
            'Validation status in CAR-T trials',
        ],
        'status': '[ ] CRITICAL - HIGHEST PRIORITY',
    },
    {
        'num': 'Q7',
        'question': 'ICAM1 & VCAM1: Role in lupus pathogenesis?',
        'needed_for': 'Discussion - Target justification',
        'search_queries': [
            'PubMed: "ICAM1" OR "VCAM1" lupus OR SLE vasculitis',
            'PubMed: "adhesion molecule" lupus pathogenesis endothelium',
            'PubMed: "ICAM1" CAR-T toxicity (endothelial injury)',
            'Target: Understand why ICAM1/VCAM1 are CAR-T targets',
        ],
        'acceptance_criteria': [
            'Elevated ICAM1/VCAM1 in SLE tissues (kidneys, skin)',
            'Role in lupus vasculitis or nephritis pathogenesis',
            'Off-target toxicity: endothelial damage risk',
            'Rationale for including in panel',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q8',
        'question': 'EGFR as CAR-T target: Is EGFR really used in CAR-T or just small molecules?',
        'needed_for': 'Methods - Target selection',
        'search_queries': [
            'PubMed: "EGFR" "CAR-T" trials',
            'PubMed: "EGFR" CAR-T cancer (find if it exists)',
            'Target: Validate EGFR as CAR-T target vs. small molecule only',
        ],
        'acceptance_criteria': [
            'Published EGFR CAR-T trials (if any)',
            'If no CAR-T trials: Decide whether to drop EGFR',
            'Rationale for including (or excluding) EGFR',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q9',
        'question': 'SenMayo VALIDATION: Has SenMayo been used in other diseases? Validated?',
        'needed_for': 'Methods - Gene panel justification',
        'search_queries': [
            'PubMed: "SenMayo" senescence signature (Saul et al. 2022)',
            'PubMed: "SenMayo" applied in [cancer, aging, autoimmunity]',
            'Target: Find 3+ papers citing SenMayo validation',
        ],
        'acceptance_criteria': [
            'Saul et al. 2022 Nature Aging paper (primary source)',
            '≥2 independent studies validating SenMayo',
            'Why SenMayo was chosen over alternatives (p16-only, etc.)',
            'Cross-platform applicability demonstrated',
        ],
        'status': '[ ] NOT STARTED',
    },
    {
        'num': 'Q10',
        'question': 'BATCH CORRECTION: Are Z-score, ComBat, MNN, Seurat compared in literature?',
        'needed_for': 'Methods - Batch correction justification',
        'search_queries': [
            'PubMed: "batch correction" "ComBat" vs "mutual nearest neighbors"',
            'PubMed: "Seurat" integration benchmarks across methods',
            'Target: Find 2-3 papers comparing batch correction methods',
        ],
        'acceptance_criteria': [
            'Published comparisons of Z-score vs. parametric methods',
            'When each method is appropriate',
            'Why Z-score chosen for 19-dataset meta-analysis',
            'Limitations of Z-score documented',
        ],
        'status': '[ ] NOT STARTED',
    },
]

for q in questions:
    doc.add_heading(f"{q['num']}: {q['question']}", level=2)

    doc.add_paragraph(f"Needed for: {q['needed_for']}", style='List Bullet')

    doc.add_paragraph('Search Queries:')
    for query in q['search_queries']:
        doc.add_paragraph(query, style='List Bullet 2')

    doc.add_paragraph('Acceptance Criteria (Question answered when ALL are true):')
    for criterion in q['acceptance_criteria']:
        doc.add_paragraph(criterion, style='List Bullet 2')

    status_para = doc.add_paragraph(q['status'])
    if 'CRITICAL' in q['status']:
        status_para.runs[0].font.bold = True
        status_para.runs[0].font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()

doc.add_page_break()

# SECTION 2: INTRODUCTION
doc.add_heading('3. SECTION 1: INTRODUCTION (MISSING - ~3-4 pages)', level=1)
doc.add_paragraph('Current: None | Needed: Full introduction with literature context')
doc.add_paragraph()

intro_subsections = [
    {
        'title': 'Subsection 1.1: SLE Epidemiology & Clinical Burden',
        'current': 'None',
        'needed': '1 page',
        'depends_on': ['Q1'],
        'outline': [
            'Global prevalence (cite: SLE affects ~5 million people)',
            'Mortality rates and organ involvement',
            'Current standard of care limitations (belimumab, steroids)',
            'Why current CAR-T (CD19) is insufficient',
        ],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 1.2: Senescence in Aging vs. Autoimmunity',
        'current': 'None',
        'needed': '1 page',
        'depends_on': ['Q3'],
        'outline': [
            'Define senescence: p16/p21, SASP, telomere shortening',
            'Senescence in healthy aging (slow accumulation)',
            'Senescence in SLE: Accelerated accumulation + inflammatory signature',
            'Why targeting senescence is uniquely relevant to SLE',
        ],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 1.3: CAR-T Immunotherapy Landscape',
        'current': 'None',
        'needed': '0.75 page',
        'depends_on': ['Q2'],
        'outline': [
            'CD19 CAR-T successes in B-cell malignancies (cite JUNO, Kymriah trials)',
            'CD19 CAR-T in SLE (cite trials if exist, else note gap)',
            'Single-target escape mechanisms (antigen loss, downregulation)',
            'Multi-target CAR-T rationale',
        ],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 1.4: Existing Senescence Markers & Research Gap',
        'current': 'None',
        'needed': '0.5 page',
        'depends_on': ['Q4', 'Q5', 'Q6'],
        'outline': [
            'Current senescence markers: p16/p21 (nuclear), SASP (soluble)',
            'Lack of surface antigen signatures on senescent cells',
            'Why SenMayo panel is comprehensive (125 genes, 7 pathways)',
            'Gap: No systematic mapping of senescence in SLE',
        ],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 1.5: This Study',
        'current': 'None',
        'needed': '0.25 page',
        'depends_on': ['Q1-Q10'],
        'outline': [
            'Hypothesis: Senescence-associated surface antigens can be identified via multi-dataset analysis',
            'Approach: Apply SenMayo to 19 GEO datasets, identify surface antigens',
            'Innovation: Multi-target CAR-T targets identified computationally',
            'Outcome: 6 candidate targets for experimental validation',
        ],
        'status': '[ ] TODO',
    },
]

for sub in intro_subsections:
    doc.add_heading(sub['title'], level=2)
    doc.add_paragraph(f"Current: {sub['current']} | Needed: {sub['needed']}")
    doc.add_paragraph(f"Depends on: {', '.join(sub['depends_on'])}")
    doc.add_paragraph('Outline:')
    for point in sub['outline']:
        doc.add_paragraph(point, style='List Bullet')
    doc.add_paragraph(sub['status'])
    doc.add_paragraph()

doc.add_page_break()

# SECTION 3: METHODS
doc.add_heading('4. SECTION 2: METHODS (PARTIAL - EXPAND)', level=1)
doc.add_paragraph('Current: docs/METHODS.md exists but needs formula detail & justifications')
doc.add_paragraph()

methods_gaps = [
    {
        'title': 'Methods 2.1: Dataset Selection Criteria (NEEDS CLARIFICATION)',
        'current': '7 inclusion / 5 exclusion criteria listed',
        'gap': 'Need to explain: Why no SLEDAI? Why datasets with n=2 patients?',
        'action': 'Add subsection: "Limitations of GEO metadata (0 datasets have SLEDAI; why we proceeded)"',
        'depends_on': ['Q1'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Methods 2.2: Batch Correction Method Choice (NEEDS JUSTIFICATION)',
        'current': 'Z-score described, but comparison missing',
        'gap': 'Need to compare: Z-score vs. ComBat vs. MNN vs. Seurat. Why NOT others?',
        'action': 'Add table comparing 4 methods + justification for Z-score choice',
        'depends_on': ['Q10'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Methods 2.3: Senescence Scoring Algorithm (NEEDS MATH)',
        'current': 'Verbal description present',
        'gap': 'Need exact formulas for Z-score, fallback logic, global re-centering',
        'action': 'Add: Mathematical formulas for S_raw, S_z, S_corrected + pseudocode',
        'depends_on': ['Q9'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Methods 2.4: Statistical Testing Details (NEEDS SPECIFICS)',
        'current': 'Tests listed (Kruskal-Wallis, Bonferroni mentioned)',
        'gap': 'Need: Exact implementation details, why non-parametric, post-hoc tests',
        'action': 'Add: Alpha levels, Bonferroni formula as implemented, normality tests if any',
        'depends_on': [],
        'status': '[ ] TODO',
    },
    {
        'title': 'Methods 2.5: CAR-T Target Selection (NEEDS LITERATURE)',
        'current': '6 targets listed with brief criteria',
        'gap': 'Need: Why these 6? Literature support for each? Why NOT others?',
        'action': 'Add: Table with clinical status of each target + citations',
        'depends_on': ['Q4', 'Q5', 'Q6', 'Q7', 'Q8'],
        'status': '[ ] TODO',
    },
]

for gap in methods_gaps:
    doc.add_heading(gap['title'], level=2)
    doc.add_paragraph(f"Current: {gap['current']}")
    doc.add_paragraph(f"Gap: {gap['gap']}")
    doc.add_paragraph(f"Action: {gap['action']}")
    if gap['depends_on']:
        doc.add_paragraph(f"Depends on: {', '.join(gap['depends_on'])}")
    doc.add_paragraph(gap['status'])
    doc.add_paragraph()

doc.add_page_break()

# SECTION 4: RESULTS
doc.add_heading('5. SECTION 3: RESULTS (GOOD - NEEDS REORGANIZATION)', level=1)
doc.add_paragraph('Current: docs/RESULTS.md has good data but conflates findings with interpretation')
doc.add_paragraph()

results_tasks = [
    {
        'title': 'Results 3.1: Remove Interpretation, Keep Pure Findings',
        'current': '"CD38, CD44 show consistent correlation with senescence"',
        'needed': '"Spearman r=0.40 (p=0.007), r=0.36 (p=0.017) in GSE228066 (n=45)"',
        'status': '[ ] TODO',
    },
    {
        'title': 'Results 3.2: Add Target Correlation Heatmap & Independence Matrix',
        'current': 'Separate correlations reported',
        'needed': 'Compute: Correlation between CD38 and CD44 expression (are they independent?)',
        'action': 'Run analysis: Correlation matrix for 6 targets in 3 datasets',
        'status': '[ ] TODO',
    },
    {
        'title': 'Results 3.3: Explain Cross-Tissue Heterogeneity (Kidney > Synovium)',
        'current': 'Variance reported but not explained',
        'needed': 'Why kidney shows senescence, synovium does not? Defer to Discussion.',
        'action': 'Move interpretation to Discussion; Results = numbers only',
        'status': '[ ] TODO',
    },
    {
        'title': 'Results 3.4: Separate Results from Discussion',
        'current': 'Findings mixed with "suggests," "indicates," "validates"',
        'needed': 'Results: data only. Discussion: interpretation',
        'action': 'Rewrite RESULTS.md as pure findings; create DISCUSSION.md for interpretation',
        'status': '[ ] TODO',
    },
]

for task in results_tasks:
    doc.add_heading(task['title'], level=2)
    doc.add_paragraph(f"Current: {task['current']}")
    doc.add_paragraph(f"Needed: {task['needed']}")
    if 'action' in task:
        doc.add_paragraph(f"Action: {task['action']}")
    doc.add_paragraph(task['status'])
    doc.add_paragraph()

doc.add_page_break()

# SECTION 5: DISCUSSION
doc.add_heading('6. SECTION 4: DISCUSSION (MISSING - ~4-5 pages)', level=1)
doc.add_paragraph('Current: None | Needed: Full discussion integrating literature')
doc.add_paragraph()

discussion_subsections = [
    {
        'title': 'Subsection 4.1: Comparison to Existing SLE Therapies',
        'needed': '1 page',
        'outline': [
            'Belimumab (B-cell activator inhibitor): Response rates, long-term outcomes',
            'CD19 CAR-T: Response rates, flare prevention, toxicity',
            'Proposed: Senescence-targeting CAR-T as complement or alternative',
            'Rationale: Why targeting senescence + CD19 may be superior to CD19 alone',
        ],
        'depends_on': ['Q1', 'Q2'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 4.2: Interpretation of Multi-Target > Single-Target',
        'needed': '0.75 page',
        'outline': [
            'Data: Combined CD38+CD44+CSPG4 (r=0.576) > individuals (r=0.23-0.40)',
            'Question: Are targets additive (different cells) or redundant (same cell)?',
            'Analysis needed: Compute correlation matrix (do targets correlate?)',
            'Mechanistic explanation: Why combination outperforms singles',
        ],
        'depends_on': ['Results 3.2'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 4.3: Off-Target Toxicity Risk Analysis',
        'needed': '0.75 page',
        'outline': [
            'CD38: Normal expression on plasma cells, T cells, NK cells',
            'CD44: Hematopoietic stem cells, memory T cells, endothelium',
            'CSPG4: [Based on literature answer to Q6]',
            'ICAM1: Endothelium (vascular injury risk)',
            'VCAM1: Endothelium, immune cells',
            'Prior CAR-T toxicities: Summarize from clinical trials',
            'Mitigation: Patient selection, combination with senolytics',
        ],
        'depends_on': ['Q4', 'Q5', 'Q6', 'Q7'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 4.4: Novelty of Each Target',
        'needed': '1 page',
        'outline': [
            'CD38: "Validated in CAR-T trials [CITE]; our contribution is validating in SLE senescence context"',
            'CD44: "Known memory marker; novel as senescence-associated surface antigen"',
            'CSPG4: "[Based on Q6 answer] If novel on immune cells: strong claim; if known: position as validation"',
            'ICAM1/VCAM1: "Endothelial markers in lupus; novel association with senescence"',
            'Overall novelty: "First comprehensive senescence-surface antigen mapping in SLE"',
        ],
        'depends_on': ['Q4', 'Q5', 'Q6'],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 4.5: Tissue-Specific Senescence Patterns',
        'needed': '0.75 page',
        'outline': [
            'Observation: Kidney senescence > synovium senescence',
            'Hypothesis: Different pathogenic mechanisms (humoral vs. cellular immunity)',
            'Literature: Lupus nephritis = antibody-driven; arthritis = mixed',
            'Implication: Senescence-targeting CAR-T may be kidney-preferential benefit',
            'Future: Tissue-specific senescence scoring in larger cohorts',
        ],
        'depends_on': [],
        'status': '[ ] TODO',
    },
    {
        'title': 'Subsection 4.6: Limitations & Caveats',
        'needed': '0.5 page',
        'outline': [
            'No SLEDAI metadata: Cannot measure senescence-activity correlation',
            'Fallback scoring: 15/19 datasets use variable genes, not SenMayo directly',
            'No functional validation: Computational discovery only; wet-lab validation required',
            'CD44/ICAM1 broad expression: Off-target toxicity risks',
            'Small sample sizes: Some datasets (n=2-3 patients)',
        ],
        'depends_on': [],
        'status': '[ ] TODO',
    },
]

for sub in discussion_subsections:
    doc.add_heading(sub['title'], level=2)
    doc.add_paragraph(f"Length: {sub['needed']}")
    doc.add_paragraph('Key Points:')
    for point in sub['outline']:
        doc.add_paragraph(point, style='List Bullet')
    doc.add_paragraph(f"Depends on: {', '.join(sub['depends_on'])}")
    doc.add_paragraph(sub['status'])
    doc.add_paragraph()

doc.add_page_break()

# SECTION 6: FUTURE DIRECTIONS
doc.add_heading('7. SECTION 5: FUTURE DIRECTIONS (MISSING - ~1-2 pages)', level=1)
doc.add_paragraph('Current: Implicit in Limitations | Needed: Explicit 3-tier roadmap')
doc.add_paragraph()

doc.add_heading('Subsection 5.1: Computational Validation (Immediate - 3-6 months)', level=2)
fd_comp = [
    '[ ] Ensembl-to-HUGO mapping for 15 fallback datasets',
    '[ ] Sensitivity analysis: Score robustness to gene selection (top 13 vs. top 20)',
    '[ ] Cross-validation: Apply SenMayo to independent SLE RNA-seq (GEO search)',
    '[ ] Benchmark: GSVA/ssGSEA vs. mean expression scoring',
]
for item in fd_comp:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Subsection 5.2: In Vitro Validation (6-12 months)', level=2)
fd_vitro = [
    '[ ] Flow cytometry: CD38/CD44/CSPG4 on senescent vs. non-senescent SLE PBMCs (n=10 active, 10 inactive)',
    '[ ] Co-culture killing: CAR-T efficacy against senescent cells (each target separately)',
    '[ ] SASP measurement: IL-6, TNF-α reduction after senescent cell clearance',
    '[ ] CSPG4 validation: Immunofluorescence on patient samples (confirm immune cell expression)',
]
for item in fd_vitro:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Subsection 5.3: In Vivo Validation (12-24 months)', level=2)
fd_vivo = [
    '[ ] Senescence burden: Quantify p16/p21 in lupus-prone mice (MRL/lpr, NZB/W)',
    '[ ] CAR-T efficacy: Multi-target vs. single-target in murine SLE',
    '[ ] Safety profile: Off-target toxicity (bone marrow, endothelium)',
    '[ ] Optimal targets: Which 1-3 targets show best efficacy/safety profile?',
]
for item in fd_vivo:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Subsection 5.4: Clinical Translation (24+ months)', level=2)
fd_clin = [
    'Patient selection criteria: High senescence burden (>median score)',
    'Combination strategies: CAR-T + senolytics (fisetin, quercetin)?',
    'Timing optimization: Early vs. late disease (when is senescence burden highest?)',
    'Flare prevention: Does senescence-targeting CAR-T reduce flare rate?',
]
for item in fd_clin:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

doc.add_page_break()

# SECTION 7: LITERATURE RESEARCH CHECKLIST
doc.add_heading('8. Literature Research Checklist', level=1)
doc.add_paragraph('Track progress of PubMed/Clinical Trials searches for each question')
doc.add_paragraph()

lit_table = doc.add_table(rows=1, cols=4)
lit_table.style = 'Light Grid Accent 1'
hdr = lit_table.rows[0].cells
hdr[0].text = 'Question'
hdr[1].text = 'Searches Completed'
hdr[2].text = 'Papers Found'
hdr[3].text = 'Key Findings'

for q_num in ['Q1', 'Q2', 'Q3', 'Q4', 'Q5', 'Q6', 'Q7', 'Q8', 'Q9', 'Q10']:
    row = lit_table.add_row()
    row.cells[0].text = q_num
    row.cells[1].text = '[ ] NOT STARTED'
    row.cells[2].text = '0'
    row.cells[3].text = '[pending]'

doc.add_paragraph()

# SECTION 8: ANALYSIS CHECKLIST
doc.add_heading('9. Analysis Checklist', level=1)
doc.add_paragraph('Track completion of computational analyses needed for manuscript')
doc.add_paragraph()

analysis_items = [
    ('Compute target correlation matrix (6 targets, 3 datasets)', '[ ] TODO'),
    ('Determine if targets are independent (Pearson r < 0.5?) or redundant', '[ ] TODO'),
    ('Verify Bonferroni correction is implemented correctly in code', '[ ] TODO'),
    ('Check: Are results reproducible if pipeline re-run?', '[ ] TODO'),
    ('Quantify: How many genes matched SenMayo per dataset? (should be in table)', '[ ] TODO'),
    ('Verify: All 19 datasets have senescence scores in external_validation/', '[ ] TODO'),
]

for item, status in analysis_items:
    doc.add_paragraph(f'{item} {status}', style='List Bullet')

doc.add_paragraph()

# SECTION 9: CROSS-VALIDATION CHECKLIST
doc.add_heading('10. Cross-Validation Checklist', level=1)
doc.add_paragraph('Ensure consistency across all sections')
doc.add_paragraph()

validation_items = [
    ('All numbers match across sections (19 datasets, 2,919 samples, 125 genes)', '[ ] TODO'),
    ('All citations formatted consistently (Author YYYY)', '[ ] TODO'),
    ('No conflation of Results and Discussion', '[ ] TODO'),
    ('Methods justify all choices (batch correction, targets, statistics)', '[ ] TODO'),
    ('Introduction flows to Methods flows to Results flows to Discussion', '[ ] TODO'),
    ('Future Directions are ambitious but achievable with resources listed', '[ ] TODO'),
    ('No overselling ("first," "revolutionary") without literature support', '[ ] TODO'),
]

for item, status in validation_items:
    doc.add_paragraph(f'{item} {status}', style='List Bullet')

doc.add_paragraph()
doc.add_page_break()

# FINAL PAGE
doc.add_heading('How to Use This Checklist', level=1)

instructions = [
    '1. PHASE 1 (Week 1): Complete all Q1-Q10 literature research. Update status from [NOT STARTED] to [COMPLETE].',
    '2. PHASE 2 (Week 2-3): Expand METHODS.md with all formulas, justifications, table comparisons.',
    '3. PHASE 3 (Week 3): Reorganize RESULTS.md: pure findings only (no interpretation).',
    '4. PHASE 4 (Week 4): Write DISCUSSION.md with literature integration.',
    '5. PHASE 5 (Week 5): Write FUTURE_DIRECTIONS.md with 3-tier roadmap.',
    '6. INTEGRATION: Merge all sections into single manuscript (Word/LaTeX).',
    '7. VALIDATION: Complete all cross-validation checklists.',
    '8. FINAL: Proofread, check citations, verify numbers match throughout.',
]

for instruction in instructions:
    doc.add_paragraph(instruction, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Target completion: 5 weeks of systematic work', style='List Bullet')
doc.add_paragraph('Blocker: Cannot write until Q1-Q10 are answered', style='List Bullet')

doc.save('manuscript_v3_checklist_June26_2026.docx')
print('[OK] Created: manuscript_v3_checklist_June26_2026.docx')
print()
print('MANUSCRIPT v3 CHECKLIST STRUCTURE:')
print('  - 10 Critical Research Questions (Q1-Q10) with acceptance criteria')
print('  - 5 Introduction subsections with dependencies')
print('  - 5 Methods expansion gaps with actions')
print('  - 4 Results reorganization tasks')
print('  - 6 Discussion subsections with outlines')
print('  - 4 Future Directions tiers (computational, in vitro, in vivo, clinical)')
print('  - Literature research tracker')
print('  - Analysis checklist')
print('  - Cross-validation checklist')
print()
print('NEXT STEP: Start with Q1-Q10 literature research (Week 1)')
